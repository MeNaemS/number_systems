from typing import Union


# The get_systems function creates a dictionary for converting numbers into Latin characters and vice versa for number
# systems from decimal to thirty hexadecimal.
def get_systems() -> dict:
    systems = []
    for system in range(36):
        systems.append((str(system), str(system)) if system < 10 else (str(system), chr(97 + (system - 10))))
    inverted_dictionary = dict([(chr(97 + (system - 10)), str(system)) for system in range(10, 36)])
    systems = dict(systems) | inverted_dictionary
    return systems


# The saving_the_translation function takes as parameters the variables path_to_dir of the string type, which stores
# the path to the directory with the saves, translation_history of the string type, which stores the text with the
# translation of numbers from one number system to another. The subroutine creates a *.docx file, if it does not exist,
# and writes the received text with translation into it, otherwise, the function opens the file and adds to it.
def saving_the_translation(path_to_dir: str, translation_history: str):
    from datetime import datetime
    from os import listdir
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    time = datetime.now()
    path_to_file = f'{path_to_dir}/date_{str(time.day)}_{str(time.month)}_{str(time.year)}.docx'
    document = Document(path_to_file) if path_to_file in listdir(path_to_dir) else Document()
    document_title = document.add_heading(f'{time.hour}:{time.minute}:{time.second}', level=2)
    document_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    recording_history = document.add_paragraph()
    recording_history.add_run(translation_history).bold = True
    document.save() if path_to_file in listdir(path_to_dir) else document.save(path_to_file)


# The encoding_tolerance function takes as parameters the variables number of string type or integer type, which stores
# the number that needs to be converted to another number system, system_1 of integer type, which stores the number
# system from which it is necessary to convert, lang of string type, which stores the language to display information.
# The subroutine checks the validity of characters for the specified number system and returns text indicating that the
# characters are invalid if they exceed the number system.
def encoding_tolerance(number: Union[int, str], system_1: int, lang: str) -> Union[str, None]:
    from data.translator import translator

    if system_1 < 10:
        for count in list(str(number)):
            if int(count) >= system_1:
                return translator(
                    russian=f'Содержатся недопустимые символы, не может быть больше {system_1}.',
                    english=f'Contains invalid characters, can not be more than {system_1}',
                    lang=lang
                )
    else:
        for count in [get_systems()[character.lower()] for character in str(number)]:
            if int(count) >= system_1:
                return translator(
                    russian=f'Содержатся недопустимые символы, не может быть больше {system_1}.\n'
                            f'Символ {get_systems()[count].upper()} по порядковому номеру является {count}.',
                    english=f'Contains invalid characters, can not be more than {system_1}\n'
                            f'The character {get_systems()[count].upper()} by ordinal number is {count}.',
                    lang=lang
                )


# The uncode_characters function takes as parameters the variables characters of a string type, which stores the
# characters whose position needs to be changed, finding a string of type, which stores the position for the
# characters. The program goes through the received characters and changes their position and returns the result.
def unicode_characters(characters: str, finding: str) -> str:
    match finding:
        case 'top':
            indexes = {"0": "\u2070", "1": "\u00B9", "2": "\u00B2", "3": "\u00B3", "4": "\u2074", "5": "\u2075",
                       "6": "\u2076", "7": "\u2077", "8": "\u2078", "9": "\u2079", "-": "\u207B"}
        case 'bottom':
            indexes = {"0": "\u2080", "1": "\u2081", "2": "\u2082", "3": "\u2083", "4": "\u2084", "5": "\u2085",
                       "6": "\u2086", "7": "\u2087", "8": "\u2088", "9": "\u2089", "-": "\u208B"}
        case _:
            raise Exception('There is no specified provision.')
    index = ''
    for element in [indexes[element] for element in list(str(characters))]:
        index += element
    return index


# The from_the_decimal function takes as parameters the variables number of lowercase or integer type, which stores the
# number that needs to be converted to another number system, system_2 of the integer type, which stores the number
# system to which it is necessary to convert. The subroutine converts the number to the specified number system and
# displays the result.
def from_the_decimal(number: Union[int, str], system_2: int) -> str:
    result = ''
    if type(number) == str: number = int(number)
    if system_2 < 10:
        while True:
            if 0 < number < system_2:
                result += str(number % system_2)
                break
            result += str(number % system_2)
            number = number // system_2
    else:
        while True:
            if number < system_2:
                if 0 < number % system_2 < 10:
                    result += str(number % system_2)
                else:
                    result += get_systems()[str(number % system_2)].upper()
                break
            if number % system_2 < 10:
                result += str(number % system_2)
            else:
                result += get_systems()[str(number % system_2)].upper()
            number = number // system_2
    result = ''.join(list(reversed(list(result))))
    return result


# The from_the_decimal function takes as parameters the variables number of lowercase or integer type, which stores the
# number that needs to be converted to another number system, system_2 of the integer type, which stores the number
# system to which it is necessary to convert. The subroutine converts the number to the specified number system and
# displays the result.
def help_from_the_decimal(number: Union[int, str], system_2: int, lang: str) -> str:
    from data.translator import translator

    result = ''
    full_answer = translator(
        russian=('\nНужно делить число и записывать остаток от деления до тех пор, пока оно не будет меньше системы, '
                 'в которую нужно перевести.\n'),
        english=('\nWe need to divide the number and write down the remainder of the division until it is less than '
                 'the system to which we need to translate.\n'),
        lang=lang
    )
    if type(number) == str: number = int(number)
    if system_2 < 10:
        while True:
            if 0 < number < system_2:
                full_answer += translator(
                    russian=f'  {number} / {system_2} = {number / system_2} (целая часть числа: {number // system_2}, '
                            f'остаток от деления: {number % system_2})\n',
                    english=f'  {number} / {system_2} = {number / system_2} (integer part of a number: '
                            f'{number // system_2}, the remainder of the division: {number % system_2})\n',
                    lang=lang
                )
                result += str(number % system_2)
                break
            full_answer += translator(
                russian=f'  {number} / {system_2} = {number / system_2} (целая часть числа: {number // system_2}, '
                        f'остаток от деления: {number % system_2})\n',
                english=f'  {number} / {system_2} = {number / system_2} (integer part of a number: '
                        f'{number // system_2}, the remainder of the division: {number % system_2})\n',
                lang=lang
            )
            result += str(number % system_2)
            number = number // system_2
    else:
        while True:
            if 0 < number < system_2:
                full_answer += translator(
                    russian=f'  {number} / {system_2} = {number / system_2} (целая часть числа: {number // system_2}, '
                            f'остаток от деления: {number % system_2})\n',
                    english=f'  {number} / {system_2} = {number / system_2} (integer part of a number: '
                            f'{number // system_2}, the remainder of the division: {number % system_2})\n',
                    lang=lang
                )
                if number % system_2 < 10:
                    result += str(number % system_2)
                else:
                    symbol = get_systems()[str(number % system_2)].upper()
                    full_answer += translator(
                        russian=f'  Затем следует перевести остаток от деления в латинский символ, который идёт '
                                f'{number % system_2} по счёту. ({number % system_2} = {symbol})\n',
                        english=f'  Then you should translate the remainder of the division into a Latin character, '
                                f'which is {number % system_2} in the count.\n',
                        lang=lang
                    )
                    result += symbol
                break
            full_answer += translator(
                russian=f'  {number} / {system_2} = {number / system_2} (целая часть числа: {number // system_2}, '
                        f'остаток от деления: {number % system_2})\n',
                english=f'  {number} / {system_2} = {number / system_2} (integer part of a number: '
                        f'{number // system_2}, the remainder of the division: {number % system_2})\n',
                lang=lang
            )
            if number % system_2 < 10:
                result += str(number % system_2)
            else:
                symbol = get_systems()[str(number % system_2)].upper()
                full_answer += translator(
                    russian=f'  Затем следует перевести остаток от деления в латинский символ, который идёт '
                            f'{number % system_2} по счёту. ({number % system_2} = {symbol})\n',
                    english=f'  Then you should translate the remainder of the division into a Latin character, which '
                            f'is {number % system_2} in the count. ({number % system_2} = {symbol})\n',
                    lang=lang
                )
                result += symbol
            number = number // system_2
    full_answer += translator(
        russian=f'В конце нужно перевернуть полученное число: {result} -> {"".join(list(reversed(list(result))))}.\n'
                f'Ответ: {"".join(list(reversed(list(result))))}',
        english=f'At the end, you need to flip the resulting number: {result} -> '
                f'{"".join(list(reversed(list(result))))}.\n'
                f'Answer: {"".join(list(reversed(list(result))))}',
        lang=lang
    )
    return full_answer


# The from_another_system function takes as parameters the variables number of a string type, which stores the number
# that needs to be converted to the decimal number system, system_1 of an integer type, which stores the initial number
# system. The subroutine converts a number from any number system to decimal and returns the resulting number.
def from_another_system(number: str, system_1: int) -> int:
    number = [get_systems()[count.lower()] for count in list(reversed(list(number)))]
    result = 0
    for count in range(len(number)):
        result += int(number[count]) * (system_1 ** count)
    return result


# The function help_from_another_system takes as parameters the variables number of a string type, which stores the
# number that needs to be converted to the decimal number system, system_1 of an integer type, which stores the initial
# number system, lang of a string type, which stores the language for displaying information. The subroutine explains
# step by step each step performed during the conversion to decimal number system and returns the resulting text with
# an explanation.
def help_from_another_system(number: str, system_1: int, lang: str) -> int:
    from data.translator import translator

    full_answer = translator(
        russian=f'\nНужно перевернуть число: {number} -> {"".join(list(reversed(list(number))))}.\n',
        english=f'\nNeed to flip the number: {number} -> {"".join(list(reversed(list(number))))}.\n',
        lang=lang
    )
    number = list(reversed(list(number)))
    for count in range(len(number)):
        if ord(number[count]) not in range(48, 58):
            full_answer += translator(
                russian=f'  Заменить латинские символы на их порядковые номера: {number[count]} -> '
                        f'{get_systems()[number[count].lower()]}.\n',
                english=f'  Replace Latin characters with their ordinal numbers: {number[count]} -> '
                        f'{get_systems()[number[count].lower()]}.\n',
                lang=lang
            )
            number[count] = get_systems()[number[count].lower()]
    result = 0
    full_answer += translator(
        russian='Затем нужно перемножить каждую цифру на её систему счисления, возведенную в порядковую степень, и '
                'сложить полученные результаты.\n',
        english='Then you need to multiply each digit by its number system, raised to an ordinal power, and add up the '
                'results obtained.\n',
        lang=lang
    )
    for count in range(len(number)):
        full_answer += (f'  {number[count]} * {system_1}{unicode_characters(characters=count, finding="top")} = '
                        f'{int(number[count]) * (system_1 ** count)}\n')
        result += int(number[count]) * (system_1 ** count)
    full_answer += translator(
        russian=f'Ответ: {result}',
        english=f'Answer: {result}',
        lang=lang
    )
    return full_answer


# The number_systems function takes as parameters the variables number of a string type, which stores the number that
# needs to be converted to another number system, system_1 of a string type, which stores the initial number system of
# the number, system_2 of a string type, which stores the number system to which it is necessary to convert , lang of a
# string type, storing the language for displaying information, full_answer of a boolean type, storing a logical value
# for outputting a detailed translation or a short one, which is False by default. The subroutine checks the tolerance
# of characters using the encoding_tolerance function, then calls the necessary function to translate the number.
def number_systems(number: str, system_1: str, system_2: str, lang: str, full_answer: bool = False) -> Union[int, str]:
    ENCODING_TOLERANCE: Union[str, None] = encoding_tolerance(number=number, system_1=int(system_1), lang=lang)
    if ENCODING_TOLERANCE is None:
        answer = ''
        if not full_answer:
            if int(system_1) != 10 and int(system_2) == 10:
                answer += str(from_another_system(number=number, system_1=int(system_1)))
            elif int(system_1) != 10 and int(system_2) != 10:
                answer += from_the_decimal(
                    number=from_another_system(number=number, system_1=int(system_1)),
                    system_2=int(system_2)
                )
            else:
                answer += from_the_decimal(number=number, system_2=int(system_2))
            return f'{answer}{unicode_characters(characters=system_2, finding="bottom")}'
        else:
            if int(system_1) != 10 and int(system_2) == 10:
                answer += help_from_another_system(number=number, system_1=int(system_1), lang=lang)
            elif int(system_1) != 10 and int(system_2) != 10:
                answer += (f'{help_from_another_system(number=number, system_1=int(system_1), lang=lang)}'
                           f'{unicode_characters(characters=10, finding="bottom")}')
                answer += help_from_the_decimal(
                    number=from_another_system(number=number, system_1=int(system_1)),
                    system_2=int(system_2),
                    lang=lang
                )
            else:
                answer += help_from_the_decimal(number=number, system_2=int(system_2), lang=lang)
            return f'{answer}{unicode_characters(characters=system_2, finding="bottom")}'
    else:
        return ENCODING_TOLERANCE
